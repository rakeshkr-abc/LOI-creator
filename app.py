import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile
import os
import subprocess
import tempfile

st.set_page_config(page_title="ABCs LOI Generator", layout="centered")

st.title("📄 ABCs LOI Generator")
st.write("Upload your Student List and Template to generate both **Word** and **PDF** files.")

# --- File Uploaders ---
col1, col2 = st.columns(2)
with col1:
    csv_file = st.file_uploader("Upload Student List (CSV/Excel)", type=['csv', 'xlsx'])
with col2:
    template_file = st.file_uploader("Upload Company LOI Template (Docx)", type=['docx'])

if csv_file and template_file:
    # Read Data
    if csv_file.name.endswith('.csv'):
        df = pd.read_csv(csv_file)
    else:
        df = pd.read_excel(csv_file)

    # Clean data (Remove empty rows)
    df = df.dropna(subset=['Student Name'])
    st.success(f"✅ Loaded {len(df)} students.")

    if st.button("🚀 Generate All Documents (Word & PDF)"):
        zip_buffer = io.BytesIO()
        
        # We need a temporary directory to handle the PDF conversion
        with tempfile.TemporaryDirectory() as tmp_dir:
            
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                progress_bar = st.progress(0)
                
                for index, row in df.iterrows():
                    # 1. Prepare Placeholder Data
                    student_name = str(row['Student Name'])
                    college_name = str(row['College Name'])
                    clean_filename = student_name.replace(' ', '_')
                    
                    # 2. Open Template
                    template_file.seek(0)
                    doc = Document(template_file)
                    
                    replacements = {
                        "<Student Name>": student_name,
                        "<College Name>": college_name
                    }
                    
                    # 3. Replace text in paragraphs and tables
                    for p in doc.paragraphs:
                        for key, val in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, val)
                    
                    for table in doc.tables:
                        for row_tab in table.rows:
                            for cell in row_tab.cells:
                                for p in cell.paragraphs:
                                    for key, val in replacements.items():
                                        if key in p.text:
                                            p.text = p.text.replace(key, val)

                    # 4. Save DOCX to the temporary directory
                    docx_path = os.path.join(tmp_dir, f"{clean_filename}.docx")
                    doc.save(docx_path)
                    
                    # 5. Convert DOCX to PDF using LibreOffice
                    # This command works on Linux servers (Streamlit Cloud / Docker)
                    try:
                        subprocess.run([
                            'soffice', 
                            '--headless', 
                            '--convert-to', 'pdf', 
                            '--outdir', tmp_dir, 
                            docx_path
                        ], check=True, capture_output=True)
                        
                        pdf_path = os.path.join(tmp_dir, f"{clean_filename}.pdf")
                        
                        # Add PDF to ZIP
                        if os.path.exists(pdf_path):
                            zip_file.write(pdf_path, f"{clean_filename}/{clean_filename}.pdf")
                        
                    except Exception as e:
                        st.error(f"Error converting {student_name} to PDF: {e}")

                    # Add DOCX to ZIP
                    zip_file.write(docx_path, f"{clean_filename}/{clean_filename}.docx")
                    
                    # Update progress
                    progress_bar.progress((index + 1) / len(df))

            # Finalize ZIP
            st.balloons()
            st.download_button(
                label="📥 Download All (ZIP: Word + PDF)",
                data=zip_buffer.getvalue(),
                file_name="Personalized_LOIs.zip",
                mime="application/zip"
            )
# import streamlit as st
# import pandas as pd
# from docx import Document
# import io
# import zipfile
# import os

# st.set_page_config(page_title="Doc Generator", layout="centered")

# st.title("📄 ABCs LOI Generator")
# st.write("Upload your Excel/CSV and Word Template to generate individual LOI files.")

# # --- File Uploaders ---
# col1, col2 = st.columns(2)
# with col1:
#     csv_file = st.file_uploader("Upload Student List CSV/Excel", type=['csv', 'xlsx'])
# with col2:
#     template_file = st.file_uploader("Upload Template of Company LOI in MS Word format", type=['docx'])

# if csv_file and template_file:
#     # Read Data
#     if csv_file.name.endswith('.csv'):
#         df = pd.read_csv(csv_file)
#     else:
#         df = pd.read_excel(csv_file)

#     # Clean data
#     df = df.dropna(subset=['Student Name'])
    
#     st.success(f"Loaded {len(df)} students.")

#     if st.button("🚀 Generate Documents"):
#         zip_buffer = io.BytesIO()
        
#         with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
#             for index, row in df.iterrows():
#                 # Load template from memory
#                 template_file.seek(0)
#                 doc = Document(template_file)
                
#                 replacements = {
#                     "<Student Name>": str(row['Student Name']),
#                     "<College Name>": str(row['College Name'])
#                 }
                
#                 # Replacement Logic
#                 for p in doc.paragraphs:
#                     for placeholder, value in replacements.items():
#                         if placeholder in p.text:
#                             p.text = p.text.replace(placeholder, value)
                
#                 for table in doc.tables:
#                     for row_tab in table.rows:
#                         for cell in row_tab.cells:
#                             for p in cell.paragraphs:
#                                 for placeholder, value in replacements.items():
#                                     if placeholder in p.text:
#                                         p.text = p.text.replace(placeholder, value)

#                 # Save doc to memory buffer
#                 doc_io = io.BytesIO()
#                 doc.save(doc_io)
#                 doc_io.seek(0)
                
#                 # Add to ZIP
#                 file_name = f"{str(row['Student Name']).replace(' ', '_')}.docx"
#                 zip_file.writestr(file_name, doc_io.getvalue())

#         st.balloons()
#         st.download_button(
#             label="📥 Download All (ZIP)",
#             data=zip_buffer.getvalue(),
#             file_name="Personalized_Documents.zip",
#             mime="application/zip"

#         )

